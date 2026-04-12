import os
import re
import logging
import uuid
from urllib.parse import urljoin
from flask import Flask, request, jsonify, render_template, session, redirect
from docx import Document
from docx.shared import Pt
from mistralai import Mistral
from flask_cors import CORS
from flask import send_from_directory
import requests
import json
import time
import traceback
from dotenv import load_dotenv
import asyncio
import subprocess
import threading
import json
import time
import asyncio
from dotenv import load_dotenv
import base64
import io
import json
import asyncio
import os
import time
import tempfile
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from mistralai import Mistral
from PIL import Image, ImageDraw, ImageFont
import imageio
import numpy as np
from mcp import ClientSession
from mcp.client.sse import sse_client

# Setup logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

load_dotenv()

app = Flask(__name__)
app.secret_key = os.urandom(24)  # Random secret key for session security
CORS(app)  # allow all origins

# Configuration
DEFAULT_JIRA_URL = os.getenv('JIRA_URL', 'https://qeemascrum.atlassian.net')
XRAY_CLOUD_BASE_URL = "https://xray.cloud.getxray.app"
XRAY_CLIENT_ID = os.getenv('XRAY_CLIENT_ID', '3DD691D61824422AAFC685BF28345BB1')
XRAY_CLIENT_SECRET = os.getenv('XRAY_CLIENT_SECRET', '67db40a621521094896a0b5839b917d27d461f5ab6b34b7b02041b87c2008896')
MCB_SERVER_URL = os.getenv('MCB_SERVER_URL', 'http://localhost:5678')
MOCK_MCB = os.getenv('MOCK_MCB', 'false').lower() == 'true'

PLAYWRIGHT_RUNS = {}

PLAYWRIGHT_CMD = os.getenv('PLAYWRIGHT_CMD', 'npx playwright test')
PLAYWRIGHT_WORKDIR = os.getenv('PLAYWRIGHT_WORKDIR') or None
MCP_SERVER_URL = "https://playwright-mcp-s8mf.onrender.com/sse"

MCP_URL       = "https://playwright-mcp-s8mf.onrender.com/sse"
MISTRAL_KEY   = "jb0XySiEnvm0r7R3HwSAWvp0aIi80K1v" #"dGA15ms96fjMa6vGmXO7veWlAKInqWVm"
MISTRAL_MODEL = "codestral-latest"

MAX_TOOL_RESULT_CHARS = 3000   # truncate huge snapshots to save tokens



MAX_TOOL_RESULT_CHARS = 3000
VIDEO_DIR = "tmp"
os.makedirs(VIDEO_DIR, exist_ok=True)


def mistral_call_with_retry(client, **kwargs):
    """Call Mistral with exponential backoff on rate limit errors."""
    for attempt in range(5):
        try:
            return client.chat.complete(**kwargs)
        except Exception as e:
            if "rate_limited" in str(e) or "1300" in str(e):
                wait = 2 ** attempt
                print(f"Rate limited — retrying in {wait}s...")
                time.sleep(wait)
            else:
                raise
    raise RuntimeError("Mistral rate limit exceeded after retries.")


async def take_screenshot(session, step_num):
    try:
        shot = await session.call_tool("browser_take_screenshot", {"type": "png"})
        for block in shot.content:
            if type(block).__name__ == "ImageContent" and hasattr(block, "data") and block.data:
                return block.data
    except Exception as e:
        print(f"  ✗ screenshot failed: {e}")
    return None
    


def decode_frame(b64_data, step):
    """Decode base64 PNG into a numpy frame using Pillow."""
    try:
        raw = base64.b64decode(b64_data)
        img = Image.open(io.BytesIO(raw)).convert("RGB")
        img = img.resize((1280, 720), Image.LANCZOS)

        # Draw label banner
        draw = ImageDraw.Draw(img)
        draw.rectangle([(0, 0), (1280, 48)], fill=(20, 20, 20))
        label = f"Step {step['step_num']}  |  {step['name']}  |  {json.dumps(step['args'])}"[:110]
        draw.text((12, 14), label, fill=(100, 230, 150))

        return np.array(img)
    except Exception as e:
        print(f"  ✗ decode_frame error: {e}")
        return None


def build_video(steps, video_path, fps=24):
    frames = []
    for step in steps:
        if not step.get("screenshot"):
            print(f"[video] step {step['step_num']} has no screenshot, skipping")
            continue
        frame = decode_frame(step["screenshot"], step)
        if frame is None:
            continue
        # Hold each step for 2 seconds
        for _ in range(3 * fps):
            frames.append(frame)
        print(f"[video] step {step['step_num']} frame added {frame.shape}")

    print(f"[video] total frames: {len(frames)}")
    if not frames:
        return False

    writer = imageio.get_writer(
        video_path,
        fps=fps,
        codec="libx264",
        pixelformat="yuv420p",   # broadest player compatibility
        macro_block_size=16,
        ffmpeg_params=[
            "-preset", "fast",
            "-crf", "23",
            "-movflags", "+faststart",  # streaming-friendly
        ],
    )
    for frame in frames:
        writer.append_data(frame)
    writer.close()


    size = os.path.getsize(video_path)
    print(f"[video] written {video_path} ({size} bytes)")
    return size > 1000


async def clear_browser_session(session):
    """Clear all auth state and start fresh login flow."""
    try:
        # Navigate to about:blank first
        await session.call_tool("browser_navigate", {"url": "about:blank"})
        print("[session] navigated to about:blank")
    except Exception as e:
        print(f"[session] navigate blank failed: {e}")
    
    try:
        # Clear ALL storage including IndexedDB, ServiceWorkers, etc
        await session.call_tool("browser_evaluate", {
            "function": """() => {
                localStorage.clear();
                sessionStorage.clear();
                // Clear all cookies
                document.cookie.split(";").forEach(c => {
                    document.cookie = c.split("=")[0].trim() + "=; expires=Thu, 01 Jan 1970 00:00:00 UTC; path=/;";
                });
            }"""
        })
        await session.call_tool("browser_evaluate", {
            "function": """() => {
                if (window.indexedDB) {
                    const dbs = await indexedDB.databases();
                    dbs.forEach(db => indexedDB.deleteDatabase(db.name));
                }
            }"""
        })
        print("[session] all storage and cookies cleared")
    except Exception as e:
        print(f"[session] clear storage failed: {e}")
    

async def run_agent(messages):
    client = Mistral(api_key=MISTRAL_KEY)
    steps  = []

    async with sse_client(MCP_URL) as (read, write):
        async with ClientSession(read, write) as session:
            await session.initialize()
            # Hard reset before every run
            await clear_browser_session(session)
            mcp_tools = await session.list_tools()
            tools = [
                {
                    "type": "function",
                    "function": {
                        "name": t.name,
                        "description": t.description or "",
                        "parameters": t.inputSchema or {},
                    },
                }
                for t in mcp_tools.tools
            ]

            for _ in range(10):
                response = mistral_call_with_retry(
                    client,
                    model=MISTRAL_MODEL,
                    messages=messages,
                    tools=tools,
                    tool_choice="auto",
                )
                msg = response.choices[0].message

                if not msg.tool_calls:
                    messages.append({"role": "assistant", "content": msg.content or ""})
                    return msg.content, steps

                messages.append({
                    "role": "assistant",
                    "content": msg.content or "",
                    "tool_calls": [
                        {
                            "id": tc.id,
                            "type": "function",
                            "function": {
                                "name": tc.function.name,
                                "arguments": tc.function.arguments,
                            },
                        }
                        for tc in msg.tool_calls
                    ],
                })

                for tc in msg.tool_calls:
                    args = (
                        json.loads(tc.function.arguments)
                        if isinstance(tc.function.arguments, str)
                        else tc.function.arguments
                    )

                    result = await session.call_tool(tc.function.name, args)
                    result_text = "\n".join(
                        b.text if hasattr(b, "text") else str(b)
                        for b in result.content
                    )
                    if len(result_text) > MAX_TOOL_RESULT_CHARS:
                        result_text = result_text[:MAX_TOOL_RESULT_CHARS] + "\n...[truncated]"

                    # Take screenshot
                    screenshot_b64 = await take_screenshot(session, len(steps) + 1)

                    steps.append({
                        "step_num":   len(steps) + 1,
                        "name":       tc.function.name,
                        "args":       args,
                        "result":     result_text,
                        "screenshot": screenshot_b64,
                    })

                    messages.append({
                        "role":         "tool",
                        "tool_call_id": tc.id,
                        "name":         tc.function.name,
                        "content":      result_text,
                    })

            return "Max turns reached.", steps


@app.post("/execute_mistral")
def execute_mistral():
    body     = request.get_json()
    messages = body.get("messages", [])

    if not messages:
        return jsonify({"error": "messages required"}), 400
    if not MISTRAL_KEY:
        return jsonify({"error": "MISTRAL_API_KEY not set"}), 500

    answer, steps = asyncio.run(run_agent(messages))

    # Build video
    video_name = f"agent_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp4"
    video_path = os.path.join(VIDEO_DIR, video_name)
    has_video  = build_video(steps, video_path)
    print(has_video)
    print(f"{steps}")
    # Strip screenshots from JSON response (they're in the video)
    clean_steps = [
        {"step": s["step_num"], "tool": s["name"], "args": s["args"], "result": s["result"]}
        for s in steps
    ]

    return jsonify({
        "answer": answer,
        "steps":  clean_steps,
        "video":  f"/tmp/{video_name}" if has_video else None,
    })


@app.get("/video/<filename>")
def get_video(filename):
    path = os.path.join(VIDEO_DIR, filename)
    if not os.path.exists(path):
        return jsonify({"error": "not found"}), 404
    return send_file(path, mimetype="video/mp4")

def start_playwright_run(spec_path):
    run_id = uuid.uuid4().hex
    PLAYWRIGHT_RUNS[run_id] = {
        "status": "starting",
        "logs": [],
        "returncode": None
    }
    def worker():
        cmd = f'{PLAYWRIGHT_CMD} "{spec_path}" --headed'
        PLAYWRIGHT_RUNS[run_id]["status"] = "running"
        try:
            proc = subprocess.Popen(
                cmd,
                shell=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                cwd=PLAYWRIGHT_WORKDIR
            )
        except Exception as e:
            PLAYWRIGHT_RUNS[run_id]["status"] = "failed"
            PLAYWRIGHT_RUNS[run_id]["logs"].append(str(e))
            return
        for line in proc.stdout:
            PLAYWRIGHT_RUNS[run_id]["logs"].append(line.rstrip())
        proc.wait()
        PLAYWRIGHT_RUNS[run_id]["returncode"] = proc.returncode
        if proc.returncode == 0:
            PLAYWRIGHT_RUNS[run_id]["status"] = "completed"
        else:
            PLAYWRIGHT_RUNS[run_id]["status"] = "failed"
    thread = threading.Thread(target=worker, daemon=True)
    thread.start()
    return run_id

@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory("/tmp", filename, as_attachment=True)

def call_mistral_model(prompt):
    try:
        client = Mistral(api_key="jb0XySiEnvm0r7R3HwSAWvp0aIi80K1v")
        model = "mistral-large-2512"
        chat_response = client.chat.complete(
            model=model,
            messages=[{"role": "user", "content": prompt}]
        )
        if chat_response:
            return chat_response.choices[0].message.content.strip()
        else:
            raise Exception("Error calling Mistral model: No response received")
    except Exception as e:
        logger.error(f"Mistral API error: {str(e)}")
        raise

def extract_sections_from_template(template_path):
    try:
        document = Document(template_path)
        sections = {}
        section_title = None
        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                section_title = paragraph.text.strip()
                sections[section_title] = ""
            elif section_title:
                sections[section_title] += paragraph.text.strip() + " "
        return sections
    except Exception as e:
        logger.error(f"Error extracting sections from template: {str(e)}")
        raise

def fill_sections_with_scope(sections, job_scope):
    filled_sections = {}
    for section, description in sections.items():
        if description.strip():
            prompt = f"Based on the job scope: {job_scope}, fill the section '{section}' with relevant content. Description: {description.strip()}"
            filled_sections[section] = call_mistral_model(prompt)
        else:
            filled_sections[section] = "No description provided."
    return filled_sections

def insert_markdown(paragraph, markdown_text):
    """
    Handles basic Markdown: bold (**text**), italic (*text*), and headings (#, ##, ###).
    """
    # Clear existing text
    paragraph.clear()

    # Otherwise, parse inline bold/italic
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', markdown_text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token.strip("*"))
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token.strip("*"))
            run.italic = True
        else:
            paragraph.add_run(token)

def update_template_with_filled_sections(template_path, filled_sections, output_path):
    try:
        document = Document(template_path)
        section_title = None

        for paragraph in document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                section_title = paragraph.text.strip()
            elif section_title and section_title in filled_sections:
                # Clear paragraph
                paragraph.text = ""
                # Insert formatted Markdown
                insert_markdown(paragraph, filled_sections[section_title])
                section_title = None

        document.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error updating template: {str(e)}")
        raise

# Xray Cloud Authentication
def get_xray_auth_token():
    """Get authentication token from Xray Cloud"""
    auth_url = f"{XRAY_CLOUD_BASE_URL}/api/v2/authenticate"
    auth_data = {
        "client_id": XRAY_CLIENT_ID,
        "client_secret": XRAY_CLIENT_SECRET
    }
    attempts = 3
    timeout_seconds = 60  # bumped to reduce spurious timeouts

    for attempt in range(1, attempts + 1):
        try:
            logger.info(f"Getting Xray auth token (attempt {attempt}/{attempts}) from: {auth_url}")
            response = requests.post(auth_url, json=auth_data, timeout=timeout_seconds)
            print(response.json())
            response.raise_for_status()

            token = response.json().strip('"')
            logger.info("Successfully obtained Xray auth token")
            
            print(f'@###############################################################{token}')
            return token

        except requests.exceptions.RequestException as e:
            logger.error(f"Xray authentication attempt {attempt} failed: {str(e)}")
            if hasattr(e, 'response') and e.response is not None:
                logger.error(f"Response Status: {e.response.status_code}")
                logger.error(f"Response Text: {e.response.text}")

            if attempt < attempts:
                backoff = 5 * attempt
                logger.info(f"Retrying in {backoff}s...")
                time.sleep(backoff)
            else:
                raise Exception(f"Xray authentication failed after {attempts} attempts: {str(e)}")

def poll_import_job_status(job_id, xray_token, max_attempts=30, delay_seconds=5):
    """Poll Xray import job status until completion with correct API endpoints"""
    headers = {
        "Authorization": f"Bearer {xray_token}",
        "Content-Type": "application/json"
    }
    
    # Try multiple possible endpoints for Xray Cloud
    endpoints = [
        f"{XRAY_CLOUD_BASE_URL}/api/v2/import/test/bulk/{job_id}",
        f"{XRAY_CLOUD_BASE_URL}/api/v1/import/test/bulk/{job_id}/status",
    ]
    
    for attempt in range(max_attempts):
        for endpoint in endpoints:
            try:
                logger.info(f"Polling attempt {attempt + 1}/{max_attempts} - Endpoint: {endpoint}")
                response = requests.get(endpoint, headers=headers, timeout=30)
                
                logger.info(f"Response status: {response.status_code}")
                
                if response.status_code == 200:
                    job_data = response.json()
                    logger.info(f"Raw job data: {json.dumps(job_data, indent=2)}")
                    
                    # Handle different response structures
                    status = job_data.get('status') or job_data.get('state') or 'UNKNOWN'
                    message = job_data.get('message', '')
                    errors = job_data.get('errors', [])
                    
                    logger.info(f"Job {job_id} status: {status}, message: {message}")
                    
                    if status.upper() in ['COMPLETED', 'DONE', 'SUCCESSFUL']:
                        logger.info("✅ Xray import job completed successfully")
                        return True, job_data, None
                    elif status.upper() in ['FAILED', 'ERROR']:
                        error_details = errors if errors else message
                        error_msg = f"Import failed: {error_details}"
                        logger.error(f"❌ Xray import job failed: {error_msg}")
                        return False, job_data, error_msg
                    elif status.upper() in ['ABORTED', 'CANCELLED']:
                        logger.error("❌ Xray import job was aborted")
                        return False, job_data, "Import job was aborted"
                    elif status.upper() in ['PROCESSING', 'IN_PROGRESS', 'PENDING']:
                        logger.info(f"⏳ Job still processing... ({status})")
                        break  # Continue to next attempt after delay
                    else:
                        logger.info(f"ℹ️ Job status: {status}, continuing to poll...")
                        break
                        
                elif response.status_code == 404:
                    logger.warning(f"Job {job_id} not found at {endpoint}")
                    continue  # Try next endpoint
                else:
                    logger.warning(f"Unexpected status {response.status_code} for {endpoint}: {response.text}")
                    continue  # Try next endpoint
                    
            except requests.exceptions.RequestException as e:
                logger.warning(f"Request error for {endpoint}: {str(e)}")
                continue  # Try next endpoint
            except Exception as e:
                logger.error(f"Unexpected error polling {endpoint}: {str(e)}")
                continue  # Try next endpoint
        
        # If we've exhausted all endpoints without a conclusive status, wait and retry
        if attempt < max_attempts - 1:
            logger.info(f"Waiting {delay_seconds} seconds before next polling attempt...")
            time.sleep(delay_seconds)
    
    timeout_msg = f"Job status polling timed out after {max_attempts} attempts. Job ID: {job_id}. Check Xray dashboard manually."
    logger.error(timeout_msg)
    return False, {"status": "TIMEOUT", "job_id": job_id}, timeout_msg  

def get_jira_issue_summary(issue_key, jira_config):
    """Fetches the summary (title) of a Jira issue."""
    try:
        base_url = jira_config['url'].rstrip('/')
        url = f"{base_url}/rest/api/3/issue/{issue_key}?fields=summary"
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json"}
        
        resp = requests.get(url, auth=auth, headers=headers, timeout=10)
        if resp.status_code == 200:
            return resp.json().get('fields', {}).get('summary')
        return None
    except Exception as e:
        logger.error(f"Error fetching jira summary: {e}")
        return None

def create_jira_issue(project_key, summary, issuetype="Test Set", jira_config=None):
    """Creates a new Jira issue (e.g., Test Set)."""
    try:
        base_url = jira_config['url'].rstrip('/')
        url = f"{base_url}/rest/api/3/issue"
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        
        payload = {
            "fields": {
                "project": {"key": project_key},
                "summary": summary,
                "issuetype": {"name": issuetype}
            }
        }
        
        resp = requests.post(url, json=payload, auth=auth, headers=headers, timeout=10)
        if resp.status_code == 201:
            return resp.json().get('key')
        else:
            logger.error(f"Failed to create jira issue: {resp.status_code} - {resp.text}")
            return None
    except Exception as e:
        logger.error(f"Error creating jira issue: {e}")
        return None

def import_tests_to_xray_bulk(test_cases, project_key, issue_key=None, test_set_key=None, poll_job=True):
    """
    Import test cases to Xray Cloud using bulk API and optionally poll job status.
    If issue_key is provided, links the created tests to that issue (e.g. User Story).
    If test_set_key is provided, adds the tests to that Test Set.
    """
    try:
        logger.info(f"=== STARTING XRAY IMPORT ===")
        logger.info(f"Project: {project_key}, Test cases: {len(test_cases)}, Linked Issue: {issue_key}, Test Set: {test_set_key}")
        
        # Get authentication token
        auth_token = get_xray_auth_token()
        if not auth_token:
            return False, None, "Failed to obtain Xray authentication token"
            
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json",
            "Accept": "application/json"
        }

        # Build test cases payload with CORRECT Xray Cloud API v2 format
        xray_tests = []
        for i, tc in enumerate(test_cases):
            logger.info(f"Processing test case {i+1}: {tc.get('summary', 'No summary')}")
            
            # Process steps
            steps_array = []
            steps_in = tc.get('steps', '')
            expected = tc.get('expected', '')
            
            if isinstance(steps_in, str) and steps_in.strip():
                # Split steps by newlines and clean them
                step_lines = [line.strip() for line in steps_in.split('\n') if line.strip()]
                expected_lines = [line.strip() for line in expected.split('\n')] if expected else []
                
                for step_idx, step_line in enumerate(step_lines):
                    step_data = {
                        "action": step_line,  # Remove numbering
                        "data": "",
                        "result": ""
                    }
                    
                    # Add expected result if available
                    if step_idx < len(expected_lines) and expected_lines[step_idx]:
                        step_data["result"] = expected_lines[step_idx]
                    elif expected and step_idx == 0:
                        step_data["result"] = expected
                    else:
                        step_data["result"] = "Verify expected behavior"
                    
                    steps_array.append(step_data)
            
            # If no steps provided, create a default step
            if not steps_array:
                steps_array = [{
                    "action": "Execute the test scenario",
                    "data": "",
                    "result": expected if expected else "Verify expected behavior"
                }]

            # CORRECTED: Build test case payload with proper test type format
            fields = {
                "project": {
                    "key": project_key
                },
                "summary": tc.get('summary', f'Test Case {i+1}')[:255],
                "description": tc.get('description', 'No description provided'),
                "issuetype": {
                    "name": "Test Case"  # This should match your Test issue type in Jira
                },
                "customfield_11604": [{"id": "11084"}],
                "customfield_12096": {
                    "id": "11449"  # This should match your Test issue type in Jira
                }
            }

            xray_test = {
                "xray_id": test_set_key,
                "xray_test_sets": [test_set_key],
                "fields": fields,
                # CORRECTED: Test type should be at the root level, not inside fields
                "testtype": "Manual",
                "steps": steps_array
            }
            
            # Add optional fields
            if tc.get('preconditions'):
                xray_test["preconditions"] = tc.get('preconditions')
            
            xray_tests.append(xray_test)

        # Prepare the final payload
        payload = xray_tests
        
        import_url = f"{XRAY_CLOUD_BASE_URL}/api/v2/import/test/bulk"

        logger.info(f"=== XRAY IMPORT PAYLOAD ===")
        logger.info(f"Import URL: {import_url}")
        logger.info(f"Number of tests: {len(xray_tests)}")
        logger.info(f"Project key: {project_key}")
        logger.info(f"Full payload: {json.dumps(payload, indent=2)}")

        # Make the API request
        logger.info("Sending request to Xray Cloud API...")
        response = requests.post(import_url, json=payload, headers=headers, timeout=60)
        
        logger.info(f"=== XRAY API RESPONSE ===")
        logger.info(f"Status Code: {response.status_code}")
        logger.info(f"Response Headers: {dict(response.headers)}")
        logger.info(f"Response Text: {response.text}")

        # Handle different response scenarios
        if response.status_code == 200:
            try:
                result = response.json()
                logger.info(f"Success response: {json.dumps(result, indent=2)}")
                
                # Check if this is a synchronous response (no job ID)
                if isinstance(result, list) and result:
                    # Synchronous success - tests were created immediately
                    created_count = len([r for r in result if r.get('status') == 'CREATED'])
                    logger.info(f"✅ Synchronous import successful - {created_count} tests created")
                    return True, {
                        "message": f"Successfully created {created_count} tests",
                        "created": created_count,
                        "total": len(test_cases),
                        "details": result
                    }, None
                else:
                    # Asynchronous response with job ID
                    job_id = result.get('jobId') or result.get('id')
                    if job_id:
                        logger.info(f"✅ Import job started - Job ID: {job_id}")
                        if poll_job:
                            return poll_import_job_status(job_id, auth_token)
                        else:
                            return True, {"jobId": job_id, "message": "Import started"}, None
                    else:
                        return True, {"message": "Import completed", "response": result}, None
                        
            except ValueError as e:
                logger.error(f"Failed to parse JSON response: {e}")
                return False, None, f"Invalid JSON response from Xray: {response.text}"

        elif response.status_code == 202:
            # Accepted for processing
            try:
                result = response.json()
                job_id = result.get('jobId') or result.get('id')
                logger.info(f"✅ Import accepted - Job ID: {job_id}")
                
                if poll_job and job_id:
                    return poll_import_job_status(job_id, auth_token)
                else:
                    return True, {"jobId": job_id, "message": "Import accepted for processing"}, None
                    
            except ValueError as e:
                return False, None, f"Invalid JSON response for 202: {response.text}"

        elif response.status_code == 400:
            error_msg = f"Bad Request: {response.text}"
            logger.error(f"❌ {error_msg}")
            return False, None, error_msg
            
        elif response.status_code == 401:
            error_msg = "Authentication failed - check Xray client credentials"
            logger.error(f"❌ {error_msg}")
            return False, None, error_msg
            
        elif response.status_code == 403:
            error_msg = "Permission denied - check project permissions"
            logger.error(f"❌ {error_msg}")
            return False, None, error_msg
            
        else:
            error_msg = f"Xray API returned status {response.status_code}: {response.text}"
            logger.error(f"❌ {error_msg}")
            return False, None, error_msg

    except requests.exceptions.RequestException as e:
        error_msg = f"Network error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        return False, None, error_msg
        
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        logger.error(f"❌ {error_msg}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return False, None, error_msg

# Jira Integration Functions
def get_user_jira_config():
    """Get Jira configuration from session or request"""
    jira_config = session.get('jira_config')
    
    if not jira_config and request.json:
        jira_config = {
            'url': request.json.get('jira_url', DEFAULT_JIRA_URL),
            'email': request.json.get('jira_email'),
            'api_token': request.json.get('jira_api_token')
        }
    
    return jira_config

def validate_jira_config(jira_config):
    """Validate Jira configuration"""
    if not jira_config:
        return False, "Jira configuration not provided"
    
    missing = []
    if not jira_config.get('url'):
        missing.append("Jira URL")
    if not jira_config.get('email'):
        missing.append("Jira Email")
    if not jira_config.get('api_token'):
        missing.append("Jira API Token")
    
    if missing:
        return False, f"Missing: {', '.join(missing)}"
    
    return True, "Valid"

def test_jira_connection(jira_config):
    """Test connection to Jira with provided credentials"""
    url = f"{jira_config['url']}/rest/api/3/myself"
    auth = (jira_config['email'], jira_config['api_token'])
    
    headers = {
        "Accept": "application/json"
    }
    
    try:
        logger.info(f"Testing Jira connection to: {url}")
        response = requests.get(url, headers=headers, auth=auth, timeout=30)
        if response.status_code == 200:
            user_data = response.json()
            logger.info(f"Jira connection successful. User: {user_data.get('displayName')}")
            return True, "Connection successful", user_data
        else:
            logger.error(f"Jira connection failed. Status: {response.status_code}, Response: {response.text}")
            return False, f"Jira API returned status {response.status_code}: {response.text}", None
    except requests.exceptions.RequestException as e:
        logger.error(f"Jira connection test failed: {str(e)}")
        return False, f"Connection failed: {str(e)}", None

def get_jira_projects(jira_config):
    """Get list of projects from Jira"""
    url = f"{jira_config['url']}/rest/api/3/project/search"
    
    auth = (jira_config['email'], jira_config['api_token'])
    
    headers = {
        "Accept": "application/json"
    }
    
    try:
        logger.info(f"Fetching Jira projects from: {url}")
        project_list = []
        start_at = 0
        max_results = 50
        
        while True:
            params = {
                'startAt': start_at,
                'maxResults': max_results
            }
            response = requests.get(url, headers=headers, auth=auth, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            projects = data.get('values', [])
            if not projects:
                break
                
            for project in projects:
                project_list.append({
                    'id': project['id'],
                    'key': project['key'],
                    'name': project['name'],
                    'projectTypeKey': project.get('projectTypeKey', 'software')
                })
            print(project_list)
                
            if data.get('isLast', True):
                break
            
            start_at += len(projects)
        
        logger.info(f"Successfully fetched {len(project_list)} projects from Jira")
        return True, project_list, None
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Jira API Error: {str(e)}")
        error_msg = f"Failed to fetch projects: {str(e)}"
        if hasattr(e, 'response') and e.response is not None:
            error_msg += f" (Status: {e.response.status_code})"
            try:
                error_detail = e.response.json()
                error_msg += f" - {error_detail}"
            except:
                error_msg += f" - {e.response.text}"
        return False, [], error_msg
    except Exception as e:
        logger.error(f"Unexpected error fetching Jira projects: {str(e)}")
        return False, [], f"Unexpected error: {str(e)}"

def get_jira_issue_types(jira_config, project_key=None):
    """Get available issue types for a project"""
    url = f"{jira_config['url']}/rest/api/3/issuetype"
    
    auth = (jira_config['email'], jira_config['api_token'])
    
    headers = {
        "Accept": "application/json"
    }
    
    try:
        logger.info(f"Fetching issue types from: {url}")
        response = requests.get(url, headers=headers, auth=auth, timeout=30)
        response.raise_for_status()
        issue_types = response.json()
        
        available_issue_types = []
        for issue_type in issue_types:
            issue_data = {
                'id': issue_type['id'],
                'name': issue_type['name'],
                'description': issue_type.get('description', ''),
                'subtask': issue_type.get('subtask', False)
            }
            
            if any(keyword in issue_type.get('name', '').lower() for keyword in ['test', 'qa', 'quality']):
                issue_data['is_test_type'] = True
            else:
                issue_data['is_test_type'] = False
                
            available_issue_types.append(issue_data)
        
        available_issue_types.sort(key=lambda x: (not x['is_test_type'], x['name']))
        
        logger.info(f"Found {len(available_issue_types)} issue types")
        return True, available_issue_types, None
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Jira API Error getting issue types: {str(e)}")
        error_msg = f"Failed to fetch issue types: {str(e)}"
        if hasattr(e, 'response') and e.response is not None:
            error_msg += f" (Status: {e.response.status_code})"
            try:
                error_detail = e.response.json()
                error_msg += f" - {error_detail}"
            except:
                error_msg += f" - {e.response.text}"
        return False, [], error_msg
    except Exception as e:
        logger.error(f"Unexpected error getting issue types: {str(e)}")
        return False, [], f"Unexpected error: {str(e)}"
    
def get_project_issue_types(jira_config, project_key):
    """Get issue types available for a specific project"""
    url = f"{jira_config['url']}/rest/api/3/issue/createmeta"
    
    auth = (jira_config['email'], jira_config['api_token'])
    
    headers = {
        "Accept": "application/json"
    }
    
    params = {
        'projectKeys': project_key,
        'expand': 'projects.issuetypes'
    }
    
    try:
        logger.info(f"Fetching project issue types for: {project_key}")
        response = requests.get(url, headers=headers, auth=auth, params=params, timeout=30)
        response.raise_for_status()
        createmeta = response.json()
        
        issue_types = []
        if createmeta.get('projects'):
            for project in createmeta['projects']:
                if project['key'] == project_key:
                    for issuetype in project.get('issuetypes', []):
                        issue_types.append({
                            'id': issuetype['id'],
                            'name': issuetype['name'],
                            'description': issuetype.get('description', ''),
                            'subtask': issuetype.get('subtask', False)
                        })
        
        logger.info(f"Found {len(issue_types)} issue types for project {project_key}")
        return True, issue_types, None
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Jira API Error getting project issue types: {str(e)}")
        error_msg = f"Failed to fetch project issue types: {str(e)}"
        if hasattr(e, 'response') and e.response is not None:
            error_msg += f" (Status: {e.response.status_code})"
            try:
                error_detail = e.response.json()
                error_msg += f" - {error_detail}"
            except:
                error_msg += f" - {e.response.text}"
        return False, [], error_msg
    except Exception as e:
        logger.error(f"Unexpected error getting project issue types: {str(e)}")
        return False, [], f"Unexpected error: {str(e)}"
def adf_to_text(adf):
    try:
        if isinstance(adf, str):
            return adf
        if not isinstance(adf, dict):
            return ""
        content = adf.get("content", [])
        parts = []
        for block in content:
            if isinstance(block, dict):
                if block.get("type") == "paragraph":
                    texts = []
                    for node in block.get("content", []):
                        if isinstance(node, dict) and node.get("type") == "text":
                            texts.append(str(node.get("text", "")))
                    if texts:
                        parts.append("".join(texts))
                elif block.get("type") == "heading":
                    texts = []
                    for node in block.get("content", []):
                        if isinstance(node, dict) and node.get("type") == "text":
                            texts.append(str(node.get("text", "")))
                    if texts:
                        parts.append("".join(texts))
        return "\n".join([p for p in parts if p.strip()])
    except Exception:
        return ""

def parse_test_cases_table(table_data):
    """Parse table data into structured test cases for Xray with dynamic column mapping"""
    test_cases = []
    
    if len(table_data) < 2:
        return test_cases
    
    headers = [h.lower() for h in table_data[0]]
    rows = table_data[1:]
    
    # Map headers to field names
    col_map = {
        'id': -1,
        'summary': -1,
        'description': -1,
        'steps': -1,
        'expected': -1
    }
    
    for i, h in enumerate(headers):
        if 'id' in h: col_map['id'] = i
        elif 'summary' in h: col_map['summary'] = i
        elif 'description' in h: col_map['description'] = i
        elif 'steps' in h or 'action' in h: col_map['steps'] = i
        elif 'expected' in h or 'result' in h: col_map['expected'] = i
        
    # Fallback for Summary if not found
    if col_map['summary'] == -1 and col_map['description'] != -1:
        col_map['summary'] = col_map['description']
    
    for i, row in enumerate(rows):
        if not row: continue
        
        # Helper to safely get value
        def get_val(idx, default=''):
            if idx != -1 and idx < len(row):
                return row[idx] if row[idx] else default
            return default

        test_case = {
            'id': get_val(col_map['id'], f"TC-{i+1:03d}"),
            'summary': get_val(col_map['summary'], f"Test Case {i+1}")[:255],
            'description': get_val(col_map['description'], 'No description'),
            'steps': get_val(col_map['steps'], 'Execute test steps'),
            'expected': get_val(col_map['expected'], 'Verify expected behavior')
        }
        
        # Clean ID
        if test_case['id'].startswith('---'): continue
        
        test_cases.append(test_case)
    
    logger.info(f"Parsed {len(test_cases)} test cases for Xray import")
    return test_cases

def import_issues_to_jira(issues, project_key):
    """
    Import Bug issues to Jira using Jira API with Steps, Expected, Actual, Bug Category.
    """
    try:
        logger.info(f"=== STARTING JIRA BUG IMPORT ===")
        jira_config = session.get('jira_config')
        if not jira_config:
            return False, None, "Jira configuration not found. Please connect to Jira first."
        
        url = f"{jira_config['url'].rstrip('/')}/rest/api/3/issue/bulk"
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json", "Content-Type": "application/json"}

        # Get issue types
        success, issue_types, error_msg = get_project_issue_types(jira_config, project_key)
        if not success:
            return False, None, f"Failed to get project issue types: {error_msg}"
        
        selected_issue_type = next((it for it in issue_types if 'bug' in it['name'].lower()), None)
        if not selected_issue_type:
            return False, None, "Bug issue type not found in project"
        logger.info(f"Selected issue type: {selected_issue_type['name']} (ID: {selected_issue_type['id']})")

        # Resolve custom field IDs from create metadata (project + issue type) to avoid hard-coded IDs
        createmeta_url = f"{jira_config['url'].rstrip('/')}/rest/api/3/issue/createmeta"
        createmeta_params = {
            "projectKeys": project_key,
            "issuetypeIds": selected_issue_type['id'],
            "expand": "projects.issuetypes.fields"
        }
        logger.info("Fetching create metadata to resolve custom fields")
        createmeta_resp = requests.get(createmeta_url, headers=headers, auth=auth, params=createmeta_params, timeout=30)
        createmeta_resp.raise_for_status()
        createmeta = createmeta_resp.json()

        fields_meta = {}
        if createmeta.get("projects"):
            for project in createmeta["projects"]:
                if project["key"] == project_key:
                    for issuetype in project.get("issuetypes", []):
                        if issuetype["id"] == selected_issue_type["id"]:
                            fields_meta = issuetype.get("fields", {})
                            break

        def match_field_id(fields_meta, candidates):
            """Return the first field id whose name contains any candidate (case-insensitive)."""
            for field_id, meta in fields_meta.items():
                name = str(meta.get("name", "")).lower()
                for cand in candidates:
                    if cand in name:
                        return field_id
            return None

        steps_id = match_field_id(fields_meta, ["steps", "reproduction steps", "repro steps"])
        expected_id = match_field_id(fields_meta, ["expected", "expected result"])
        actual_id = match_field_id(fields_meta, ["actual", "actual result"])
        bug_category_id = match_field_id(fields_meta, ["bug category", "category"])

        # If custom fields are missing, we'll include them in description instead
        use_description_fallback = not steps_id or not expected_id or not actual_id
        
        if not bug_category_id:
            logger.warning("Bug Category field not found, will use default value")

        # Get valid Bug Category options (prefer create metadata so we have option IDs)
        valid_categories = []
        category_lookup = {}
        if bug_category_id in fields_meta:
            meta_field = fields_meta[bug_category_id]
            allowed = meta_field.get("allowedValues", []) or meta_field.get("options", [])
            valid_categories = allowed
            # Build lookup with multiple keys for flexible matching
            for opt in allowed:
                if isinstance(opt, dict) and opt.get("value"):
                    value = str(opt.get("value", "")).strip()
                    value_lower = value.lower()
                    # Store with lowercase key
                    category_lookup[value_lower] = opt
                    # Also store with ID if available
                    if opt.get("id"):
                        category_lookup[f"id:{opt.get('id')}"] = opt

        # Fallback: fetch options directly
        if not valid_categories and bug_category_id:
            try:
                category_url = f"{jira_config['url'].rstrip('/')}/rest/api/3/customField/{bug_category_id}/option"
                response = requests.get(category_url, auth=auth, timeout=30)
                if response.status_code == 200:
                    data = response.json()
                    valid_categories = data.get('values', data.get('options', []))
                    # Build lookup with multiple keys
                    for opt in valid_categories:
                        if isinstance(opt, dict) and opt.get("value"):
                            value = str(opt.get("value", "")).strip()
                            value_lower = value.lower()
                            category_lookup[value_lower] = opt
                            if opt.get("id"):
                                category_lookup[f"id:{opt.get('id')}"] = opt
                else:
                    logger.warning(f"Unable to fetch Bug Category options (status {response.status_code}); will skip category field")
            except Exception as e:
                logger.warning(f"Error fetching Bug Category options: {str(e)}; will skip category field")

        logger.info(f"Valid Bug Category options: {[opt.get('value') for opt in valid_categories if isinstance(opt, dict) and opt.get('value')]}")

        def to_adf(text):
            if not text:
                text = "-"
            return {
                "type": "doc",
                "version": 1,
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": line}]
                    } for line in str(text).split("\n") if line.strip()
                ]
            }

        def build_description(steps, expected, actual, description=None):
            """Build description field with steps, expected, and actual"""
            desc_parts = []
            
            if description:
                desc_parts.append(description)
                desc_parts.append("")
            
            if steps:
                desc_parts.append("*Steps to Reproduce:*")
                desc_parts.append(steps)
                desc_parts.append("")
            
            if expected:
                desc_parts.append("*Expected Behavior:*")
                desc_parts.append(expected)
                desc_parts.append("")
            
            if actual:
                desc_parts.append("*Actual Behavior:*")
                desc_parts.append(actual)
            
            return "\n".join(desc_parts) if desc_parts else "No description provided."

        issue_requests = []
        for i, issue in enumerate(issues):
            summary = issue.get('summary', f'Bug {i+1}')[:255]
            
            # Get fields
            steps = issue.get('steps', '').strip()
            expected = issue.get('expected', '').strip()
            actual = issue.get('actual', '').strip()
            description = issue.get('description', '').strip()
            bug_category_raw = issue.get('bug_category', 'Functional/Code')
            bug_category = 'Functional/Code'

            # Validate Bug Category (case-insensitive) and use option id/value
            # bug_category = None
            if bug_category_id:
                if category_lookup:
                    # Try to find exact match
                    key = str(bug_category_raw).strip().lower() if bug_category_raw else ''
                    match = category_lookup.get(key)
                    
                    if match:
                        # Use the matched option
                        if match.get("id"):
                            bug_category = {"id": str(match.get("id"))}
                        elif match.get("value"):
                            bug_category = {"value": str(match.get("value"))}
                        else:
                            logger.warning(f"Issue {i+1} Bug Category match has no id or value")
                    else:
                        # Try partial match
                        found_match = None
                        for cat_key, cat_option in category_lookup.items():
                            if key in cat_key or cat_key in key:
                                found_match = cat_option
                                break
                        
                        if found_match:
                            if found_match.get("id"):
                                bug_category = {"id": str(found_match.get("id"))}
                            elif found_match.get("value"):
                                bug_category = {"value": str(found_match.get("value"))}
                        else:
                            # Try to use first available option as default
                            if valid_categories and len(valid_categories) > 0:
                                first_option = valid_categories[0]
                                if isinstance(first_option, dict):
                                    if first_option.get("id"):
                                        bug_category = {"id": str(first_option.get("id"))}
                                    elif first_option.get("value"):
                                        bug_category = {"value": str(first_option.get("value"))}
                                logger.warning(f"Issue {i+1} Bug Category '{bug_category_raw}' not found, using default: {bug_category}")
                            else:
                                logger.warning(f"Issue {i+1} Bug Category '{bug_category_raw}' not found and no defaults available, skipping category")
                else:
                    # No category lookup available, try to use provided value directly
                    if bug_category_raw:
                        bug_category = {"value": str(bug_category_raw)}
                    else:
                        logger.warning(f"Issue {i+1} No bug category provided and no lookup available")

            # Build description with steps, expected, and actual
            full_description = build_description(steps, expected, actual, description)
            description_adf = to_adf(full_description)

            fields = {
                "project": {"key": project_key},
                "summary": summary,
                "issuetype": {"id": selected_issue_type['id']},
                "description": description_adf,
            }

            # Add custom fields if available
            if steps_id and steps:
                fields[steps_id] = to_adf(steps)
            if expected_id and expected:
                fields[expected_id] = to_adf(expected)
            if actual_id and actual:
                fields[actual_id] = to_adf(actual)
            
            # Only add bug category if we have a valid value
            if bug_category_id and bug_category:
                logger.info(f"Issue {i+1} Bug Category: {bug_category}")
                fields[bug_category_id] = [bug_category]
            elif bug_category_id:
                fields[bug_category_id] = [bug_category]
                # Bug category field exists but we don't have a valid value - skip it to avoid error
                logger.warning(f"Issue {i+1} Skipping bug category field to avoid validation error")

            # Optional priority
            if issue.get('priority'):
                priority_map = {'highest': '1','high': '2','medium': '3','low': '4','lowest': '5'}
                if issue['priority'].lower() in priority_map:
                    fields['priority'] = {'id': priority_map[issue['priority'].lower()]}

            issue_requests.append({"update": {}, "fields": fields})

        if not issue_requests:
            return False, None, "No valid issues to create. All issues missing required fields or invalid Bug Category."

        payload = {"issueUpdates": issue_requests}
        logger.info(f"Payload: {issue_requests}")
        logger.info(f"Sending {len(issue_requests)} issues to Jira API...")
        response = requests.post(url, headers=headers, auth=auth, json=payload, timeout=60)
        logger.info(f"Jira response status: {response.status_code}")

        if response.status_code in [200, 201]:
            result = response.json()
            created_issues = result.get('issues', [])
            errors = result.get('errors', [])
            return True, {"created": len(created_issues), "failed": len(errors), "issues": created_issues, "errors": errors}, None
        else:
            return False, None, f"Jira API returned {response.status_code}: {response.text}"

    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return False, None, str(e)

def parse_bug_cases_table(table_data):
    """Parse table data into structured bug cases"""
    bug_cases = []
    
    if len(table_data) < 2:
        return bug_cases
    
    headers = [h.lower() if h else '' for h in table_data[0]]
    rows = table_data[1:]
    
    for i, row in enumerate(rows):
        # Skip empty rows or separator rows
        if not any(row) or (len(row) > 0 and str(row[0]).startswith('---')):
            continue
            
        bug_case = {
            'id': f"BUG-{i+1:03d}",
            'summary': '',
            'description': '',
            'steps': '',
            'expected': '',
            'actual': '',
            'priority': 'Medium',
            'severity': 'Major',
            'bug_category': 'Functional'
        }
        
        # Map columns based on headers
        for j, header in enumerate(headers):
            if j < len(row) and row[j]:
                header_lower = header.lower()
                
                if 'summary' in header_lower or 'title' in header_lower:
                    bug_case['summary'] = row[j]
                    bug_case['description'] = row[j]
                elif 'description' in header_lower:
                    bug_case['description'] = row[j]
                elif 'steps' in header_lower or 'reproduce' in header_lower:
                    bug_case['steps'] = row[j]
                elif 'expected' in header_lower:
                    bug_case['expected'] = row[j]
                elif 'actual' in header_lower:
                    bug_case['actual'] = row[j]
                elif 'priority' in header_lower:
                    bug_case['priority'] = row[j]
                elif 'severity' in header_lower:
                    bug_case['severity'] = row[j]
                elif 'category' in header_lower:
                    bug_case['bug_category'] = row[j]
        
        # Fallback if no headers matched
        if not bug_case['summary']:
            if len(row) > 1:
                bug_case['summary'] = row[1] if row[1] else f"Bug {i+1}"
                bug_case['description'] = row[1] if row[1] else ''
            if len(row) > 2:
                bug_case['steps'] = row[2] if row[2] else ''
            if len(row) > 3:
                bug_case['expected'] = row[3] if row[3] else ''
            if len(row) > 4:
                bug_case['actual'] = row[4] if row[4] else ''
            if len(row) > 5:
                bug_case['priority'] = row[5] if row[5] else 'Medium'
            if len(row) > 6:
                bug_case['severity'] = row[6] if row[6] else 'Major'
            if len(row) > 7:
                bug_case['bug_category'] = row[7] if row[7] else 'Functional'
        
        # Ensure summary is not empty
        if not bug_case['summary']:
            bug_case['summary'] = f"Bug {i+1}"
        
        bug_cases.append(bug_case)
    
    logger.info(f"Parsed {len(bug_cases)} bug cases")
    return bug_cases

# ==================== ENDPOINTS ====================

def add_tests_to_xray_test_set(test_keys, test_set_key):
    """Adds multiple test cases to a Test Set using Xray Cloud API v2."""
    try:
        auth_token = get_xray_auth_token()
        if not auth_token:
            return False, "Failed to obtain Xray token"
            
        url = f"{XRAY_CLOUD_BASE_URL}/api/v2/testset/{test_set_key}/testcases"
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json"
        }
        payload = {"add": test_keys}
        
        logger.info(f"Adding {len(test_keys)} tests to Test Set {test_set_key}...")
        resp = requests.post(url, json=payload, headers=headers, timeout=30)
        
        if resp.status_code == 200:
            logger.info(f"✅ Successfully added tests to Test Set {test_set_key}")
            return True, None
        else:
            error_msg = f"Xray API error {resp.status_code}: {resp.text}"
            logger.error(error_msg)
            return False, error_msg
    except Exception as e:
        logger.error(f"Error adding tests to test set: {str(e)}")
        return False, str(e)

def link_jira_issues(source_keys, target_key, link_type="Tests"):
    """
    Link multiple source issues to a target issue.
    Default link type is 'Tests' (Source tests Target).
    Attempts to find a valid link type if default is missing.
    """
    try:
        jira_config = get_user_jira_config()
        if not jira_config:
            return False, "Jira configuration required"
            
        base_url = jira_config['url'].rstrip('/')
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        
        # 1. Resolve Link Type
        resolved_link_type = link_type
        try:
            link_types_url = f"{base_url}/rest/api/3/issueLinkType"
            lt_resp = requests.get(link_types_url, headers=headers, auth=auth, timeout=10)
            
            if lt_resp.status_code == 200:
                available_types = lt_resp.json().get('issueLinkTypes', [])
                logger.info(f"Available Issue Link Types: {lt_resp}")
                type_names = [t.get('name') for t in available_types]
                logger.info(f"Available Issue Link Types: {type_names}")
                
                # Check if requested type exists
                if link_type not in type_names:
                    logger.warning(f"Link type '{link_type}' not found. Searching for alternatives...")
                    
                    # Preferred alternatives in order
                    alternatives = ["Test", "Tests", "Relates", "Blocks", "Link", "Cloners"]
                    
                    found = False
                    for alt in alternatives:
                        # Case insensitive check
                        match = next((t['name'] for t in available_types if t['name'].lower() == alt.lower()), None)
                        if match:
                            resolved_link_type = match
                            logger.info(f"Using alternative link type: '{resolved_link_type}'")
                            found = True
                            break
                    
                    if not found and available_types:
                        # Fallback to the first available one if nothing matches
                        resolved_link_type = available_types[0]['name']
                        logger.warning(f"No suitable link type found. Fallback to first available: '{resolved_link_type}'")
            else:
                logger.warning(f"Failed to fetch issue link types: {lt_resp.status_code}")
                
        except Exception as e:
            logger.error(f"Error resolving link types: {e}")
            # Proceed with default if resolution fails

        # 2. Perform Linking
        link_url = f"{base_url}/rest/api/3/issueLink"
        success_count = 0
        errors = []
        
        for src_key in source_keys:
            payload = {
                "type": {
                    "name": resolved_link_type
                },
                "outwardIssue": {
                    "key": src_key
                },
                "inwardIssue": {
                    "key": target_key
                }
            }
            
            try:
                response = requests.post(link_url, headers=headers, auth=auth, json=payload, timeout=10)
                if response.status_code == 201:
                    success_count += 1
                else:
                    errors.append(f"Failed to link {src_key}: {response.status_code} - {response.text}")
                    logger.error(f"Link failed for {src_key} with type '{resolved_link_type}': {response.text}")
            except Exception as e:
                errors.append(f"Error linking {src_key}: {str(e)}")
                
        if success_count == len(source_keys):
            return True, f"Successfully linked {success_count} issues to {target_key} using '{resolved_link_type}'"
        elif success_count > 0:
            return True, f"Linked {success_count}/{len(source_keys)} issues using '{resolved_link_type}'. Errors: {'; '.join(errors[:3])}..."
        else:
            return False, f"Failed to link issues using '{resolved_link_type}': {'; '.join(errors[:3])}"
            
    except Exception as e:
        logger.error(f"Error in link_jira_issues: {str(e)}")
        return False, str(e)
        return False, f"Linking error: {str(e)}"

@app.route('/jira/import-bugs', methods=['POST'])
def import_bugs_to_xray():
    """Import bug reports to Jira"""
    try:
        data = request.json
        logger.info(f"=== BUGS IMPORT REQUEST ===")
        
        table_data = data.get('table_data', [])
        project_key = data.get('project_key')
        issue_key = data.get('issue_key')
        
        if not table_data or len(table_data) < 2:
            return jsonify({"error": "No valid table data provided"}), 400
        
        if not project_key:
            return jsonify({"error": "Project key is required"}), 400
        
        bug_cases = parse_bug_cases_table(table_data)
        logger.info(f"Parsed {len(bug_cases)} bug cases")
        
        if not bug_cases:
            return jsonify({"error": "No valid bug cases found in table data"}), 400
        
        success, result, error_msg = import_issues_to_jira(bug_cases, project_key)
        
        if success:
            if issue_key:
                created_keys = [bug.get('key') for bug in result.get('issues', []) if bug.get('key')]
                if created_keys:
                    link_success, link_msg = link_jira_issues(created_keys, issue_key, link_type="Blocks")
                    if link_success:
                        result['linked_count'] = len(created_keys)
                        result['linked_issue'] = issue_key
                    else:
                        result['linking_error'] = link_msg
            return jsonify({
                "success": True,
                "message": f"Successfully submitted {len(bug_cases)} bug reports to Jira",
                "result": result
            })
        else:
            return jsonify({
                "success": False,
                "error": error_msg
            }), 400
        
    except Exception as e:
        logger.error(f"Error in import_bugs_to_jira: {str(e)}")
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

# Jira Configuration Endpoints
@app.route('/jira/test-connection', methods=['POST'])
def test_jira_connection_endpoint():
    """Test Jira connection with user credentials"""
    try:
        data = request.json
        jira_config = {
            'url': data.get('jira_url', DEFAULT_JIRA_URL).rstrip('/'),
            'email': data.get('jira_email'),
            'api_token': data.get('jira_api_token')
        }
        
        is_valid, validation_msg = validate_jira_config(jira_config)
        if not is_valid:
            return jsonify({"success": False, "error": validation_msg}), 400
        
        success, message, user_data = test_jira_connection(jira_config)
        
        if success:
            session['jira_config'] = jira_config
            session['jira_user'] = user_data.get('displayName', 'Unknown User')
            return jsonify({
                "success": True, 
                "message": f"Connected successfully as {user_data.get('displayName')}",
                "user": user_data.get('displayName')
            })
        else:
            return jsonify({"success": False, "error": message}), 400
            
    except Exception as e:
        logger.error(f"Error testing Jira connection: {str(e)}")
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

@app.route('/jira/projects', methods=['GET', 'POST'])
def get_jira_projects_endpoint():
    """Get available Jira projects"""
    try:
        jira_config = get_user_jira_config()
        
        if request.method == 'GET':
            if not jira_config:
                return jsonify({"success": False, "error": "No Jira configuration found. Please connect first."}), 400
        else:
            data = request.json
            jira_config = {
                'url': data.get('jira_url', DEFAULT_JIRA_URL).rstrip('/'),
                'email': data.get('jira_email'),
                'api_token': data.get('jira_api_token')
            }
        
        is_valid, validation_msg = validate_jira_config(jira_config)
        if not is_valid:
            return jsonify({"success": False, "error": validation_msg}), 400
        
        success, projects, error_msg = get_jira_projects(jira_config)
        
        if success:
            if request.method == 'POST':
                session['jira_config'] = jira_config
            return jsonify({"success": True, "projects": projects})
        else:
            return jsonify({"success": False, "error": error_msg}), 400
            
    except Exception as e:
        logger.error(f"Error in get_jira_projects_endpoint: {str(e)}")
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

@app.route('/jira/issue-types', methods=['GET'])
def get_jira_issue_types_endpoint():
    """Get available issue types for a project"""
    try:
        jira_config = get_user_jira_config()
        if not jira_config:
            return jsonify({"success": False, "error": "Jira configuration required"}), 400
        
        project_key = request.args.get('projectKey')
        
        if project_key:
            success, issue_types, error_msg = get_project_issue_types(jira_config, project_key)
        else:
            success, issue_types, error_msg = get_jira_issue_types(jira_config)
        
        if success:
            return jsonify({"success": True, "issue_types": issue_types})
        else:
            return jsonify({"success": False, "error": error_msg}), 400
            
    except Exception as e:
        logger.error(f"Error in get_jira_issue_types_endpoint: {str(e)}")
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500
@app.route('/jira/issues', methods=['GET'])
def search_jira_issues():
    try:
        jira_config = get_user_jira_config()
        if not jira_config:
            return jsonify({"success": False, "error": "Jira configuration required"}), 400
        project_key = request.args.get('projectKey', '').strip()
        q = request.args.get('q', '').strip()
        if not project_key:
            return jsonify({"success": False, "error": "projectKey is required"}), 400
        search_url = f"{jira_config['url'].rstrip('/')}/rest/api/3/search/jql"
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        def looks_like_key(s):
            return bool(re.match(r'^[A-Z][A-Z0-9_]*-\d+$', s))
        jql_parts = [f'project = "{project_key}"']
        story_types = []
        try:
            st_ok, st_list, st_err = get_project_issue_types(jira_config, project_key)
            if st_ok and st_list:
                story_types = [it.get('name') for it in st_list if 'story' in str(it.get('name','')).lower()]
        except Exception:
            pass
        if story_types:
            escaped = [name.replace('"','\\"') for name in story_types]
            jql_parts.append('issuetype in ("' + '","'.join(escaped) + '")')
        else:
            jql_parts.append('issuetype not in ("Test","Precondition","Test Set","Test Execution","Sub-Test Execution")')
        if q:
            q_escaped = q.replace('"', '\\"')
            if looks_like_key(q):
                jql_parts.append(f'key = "{q_escaped}"')
            elif len(q) >= 3:
                jql_parts.append(f'(summary ~ "{q_escaped}" OR description ~ "{q_escaped}")')
        jql = " AND ".join(jql_parts) + " ORDER BY created DESC"
        search_payload = {"jql": jql, "maxResults": 50}
        search_resp = requests.post(search_url, headers=headers, auth=auth, json=search_payload, timeout=30)
        if search_resp.status_code != 200:
            return jsonify({"success": False, "error": f"Jira API returned {search_resp.status_code}: {search_resp.text}"}), 400
        search_data = search_resp.json() or {}
        raw_items = search_data.get("issues", [])
        ids_or_keys = []
        for it in raw_items:
            key = it.get("key")
            id_ = it.get("id")
            ids_or_keys.append(key or id_)
        if not ids_or_keys:
            return jsonify({"success": True, "issues": [], "count": 0})
        bulk_url = f"{jira_config['url'].rstrip('/')}/rest/api/3/issue/bulkfetch"
        bulk_payload = {
            "issueIdsOrKeys": ids_or_keys,
            "fields": ["summary", "description", "status", "issuetype"]
        }
        bulk_resp = requests.post(bulk_url, headers=headers, auth=auth, json=bulk_payload, timeout=30)
        if bulk_resp.status_code != 200:
            return jsonify({"success": False, "error": f"Bulk fetch failed {bulk_resp.status_code}: {bulk_resp.text}"}), 400
        bulk = bulk_resp.json() or {}
        result_issues = []
        for item in bulk.get("issues", []):
            fields = item.get("fields", {}) or {}
            desc = fields.get("description")
            text_desc = adf_to_text(desc) if desc is not None else ""
            result_issues.append({
                "key": item.get("key"),
                "summary": fields.get("summary"),
                "description": text_desc,
                "status": (fields.get("status") or {}).get("name"),
                "issuetype": (fields.get("issuetype") or {}).get("name")
            })
        return jsonify({"success": True, "issues": result_issues, "count": len(result_issues)})
    except Exception as e:
        logger.error(f"Error searching Jira issues: {str(e)}")
        return jsonify({"success": False, "error": f"Search failed: {str(e)}"}), 500

@app.route('/jira/disconnect', methods=['POST'])
def disconnect_jira():
    """Clear Jira configuration from session"""
    session.pop('jira_config', None)
    session.pop('jira_user', None)
    return jsonify({"success": True, "message": "Disconnected from Jira"})

@app.route('/jira/user', methods=['GET'])
def get_jira_user():
    """Get current Jira user from session"""
    jira_user = session.get('jira_user')
    jira_config = session.get('jira_config')
    
    if jira_user and jira_config:
        return jsonify({
            "connected": True,
            "user": jira_user,
            "email": jira_config.get('email'),
            "url": jira_config.get('url')
        })
    else:
        return jsonify({"connected": False})

# Xray Endpoints
@app.route('/xray/test-connection', methods=['GET'])
def test_xray_connection():
    """Test Xray Cloud connection"""
    try:
        auth_token = get_xray_auth_token()
        
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json"
        }
        
        print(headers)
        # Test basic connectivity
        test_url = f"{XRAY_CLOUD_BASE_URL}/api/v2/projects"
        print(test_url)
        try:
            response = requests.get(test_url, headers=headers, timeout=30)
            if response.status_code == 200:
                return jsonify({
                    "success": True,
                    "message": "Xray Cloud connection successful",
                    "token_obtained": True
                })
            else:
                return jsonify({
                    "success": False,
                    "error": f"Xray API returned status {response.status_code}"
                }), 500
                
        except requests.exceptions.RequestException as e:
            return jsonify({
                "success": False,
                "error": f"Xray API request failed: {str(e)}"
            }), 500
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Xray connection test failed: {str(e)}"
        }), 500

@app.route('/xray/check-job-status', methods=['POST'])
def check_xray_job_status():
    """Check the status of Xray import job"""
    try:
        data = request.json
        job_id = data.get('job_id')
        
        if not job_id:
            return jsonify({"error": "Job ID is required"}), 400
        
        auth_token = get_xray_auth_token()
        success, job_data, error_msg = poll_import_job_status(job_id, auth_token)
        
        if success:
            return jsonify({
                "success": True,
                "job_status": job_data,
                "message": f"Job completed successfully"
            })
        else:
            return jsonify({
                "success": False,
                "job_status": job_data,
                "error": error_msg or "Job failed or timed out"
            })
            
    except Exception as e:
        logger.error(f"Error checking job status: {str(e)}")
        return jsonify({"success": False, "error": f"Job status check failed: {str(e)}"}), 500

@app.route('/xray/search-tests', methods=['POST'])
def search_xray_tests():
    """Search for tests in Xray by project"""
    try:
        data = request.json
        project_key = data.get('project_key')
        
        if not project_key:
            return jsonify({"error": "Project key is required"}), 400
        
        auth_token = get_xray_auth_token()
        
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json"
        }
        
        search_url = f"{XRAY_CLOUD_BASE_URL}/api/v2/tests"
        params = {
            'projectKey': project_key,
            'limit': 50
        }
        
        response = requests.get(search_url, headers=headers, params=params, timeout=30)
        
        logger.info(f"Search tests response: {response.status_code}")
        
        if response.status_code == 200:
            tests = response.json()
            
            logger.info(f"Found {len(tests) if isinstance(tests, list) else 0} tests for project {project_key}")
            
            if isinstance(tests, list) and tests:
                formatted_tests = []
                for test in tests:
                    formatted_tests.append({
                        'id': test.get('id'),
                        'key': test.get('key'),
                        'summary': test.get('summary'),
                        'description': test.get('description', '')[:100] + '...' if test.get('description') else '',
                        'testtype': test.get('testtype', 'Manual'),
                        'projectKey': test.get('projectKey')
                    })
                
                return jsonify({
                    "success": True,
                    "message": f"✅ Found {len(tests)} tests in project {project_key}",
                    "tests": formatted_tests,
                    "count": len(tests)
                })
            else:
                return jsonify({
                    "success": True,
                    "message": f"🔍 No tests found in project {project_key}",
                    "tests": [],
                    "count": 0
                })
        else:
            return jsonify({
                "success": False,
                "error": f"Failed to search tests: {response.status_code} - {response.text}"
            })
            
    except Exception as e:
        logger.error(f"Error searching Xray tests: {str(e)}")
        return jsonify({"success": False, "error": f"Search failed: {str(e)}"}), 500

@app.route('/xray/test-sets', methods=['GET'])
def get_xray_test_sets():
    """Fetch all Test Sets for a project from Jira with auto-discovery of issue type name"""
    try:
        jira_config = get_user_jira_config()
        if not jira_config:
            return jsonify({"success": False, "error": "Jira configuration required"}), 400
            
        project_key = request.args.get('projectKey')
        if not project_key:
            return jsonify({"success": False, "error": "Project Key is required"}), 400
            
        # 1. Discover the "Test Set" issue type name for this project
        test_set_type_name = "Test Set" # Default
        try:
            success_it, issue_types, err_it = get_project_issue_types(jira_config, project_key)
            if success_it and issue_types:
                # Look for names containing "Test Set"
                found_type = next((it for it in issue_types if "test set" in (it.get('name') or "").lower()), None)
                if found_type:
                    test_set_type_name = found_type['name']
                    logger.info(f"Auto-discovered Test Set issue type name: '{test_set_type_name}'")
                else:
                    names = [it.get('name') for it in issue_types]
                    logger.warning(f"Could not find exact 'Test Set' issue type. Available: {names}")
            else:
                logger.warning(f"Failed to fetch issue types for discovery: {err_it}")
        except Exception as e:
            logger.error(f"Error during issue type discovery: {str(e)}")

        # 2. Search for issues of discovered type
        search_url = f"{jira_config['url'].rstrip('/')}/rest/api/3/search/jql"
        auth = (jira_config['email'], jira_config['api_token'])
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        
        jql = f'project = "{project_key}" AND issuetype = "{test_set_type_name}" ORDER BY created DESC'
        logger.info(f"Fetching test sets with JQL: {jql}")
        
        payload = {
            "jql": jql,
            "maxResults": 500,
            "fields": ["summary", "key"]
        }
        payload = json.dumps( {
            "expand": "names",
            "fields":["*all"],
            "fieldsByKeys": True,
            "jql": jql,
            "maxResults": 500
        } )
        
        resp = requests.post(search_url, data=payload, headers=headers, auth=auth, timeout=30)
        
        if resp.status_code == 200:
            issues = resp.json().get('issues', [])
            test_sets = [{
                'key': issue['key'],
                'summary': issue['fields']['summary']
            } for issue in issues]
            return jsonify({"success": True, "test_sets": test_sets})
        else:
            error_data = resp.text
            try:
                error_json = resp.json()
                if 'errorMessages' in error_json and error_json['errorMessages']:
                    error_data = " ".join(error_json['errorMessages'])
            except:
                pass
            
            # If specifically an issue type error, provide more context
            if "issuetype" in error_data.lower() or "issue type" in error_data.lower():
                error_data += f" (Note: Tried searching for issue type '{test_set_type_name}')"
                
            logger.error(f"Jira API error ({resp.status_code}): {error_data}")
            return jsonify({"success": False, "error": f"Jira Error ({resp.status_code}): {error_data}"}), 500
            
    except Exception as e:
        logger.error(f"Exception fetching test sets: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/xray/test-set-tests', methods=['POST'])
def get_test_set_tests():
    """Fetch all test cases belonging to selected Test Sets"""
    try:
        data = request.json
        test_set_keys = data.get('test_set_keys', [])
        
        if not test_set_keys:
            return jsonify({"success": False, "error": "No Test Set keys provided"}), 400
            
        auth_token = get_xray_auth_token()
        if not auth_token:
            return jsonify({"success": False, "error": "Failed to authenticate with Xray"}), 500
            
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json"
        }
        
        all_test_keys = set()
        test_steps_map = {} # Map key -> {steps_text, expected_text}
        test_to_set_map = {} # Map test_key -> test_set_key
        
        if test_set_keys:
            # Xray Cloud's getTestSet expects a numerical issueId. 
            # To fetch by Key (like AHC-1535), we use getTestSets with a JQL filter.
            gql_url = f"{XRAY_CLOUD_BASE_URL}/api/v2/graphql"
            
            # Escape keys for JQL
            keys_str = ", ".join([f"\"{k}\"" for k in test_set_keys])
            jql = f"key in ({keys_str})"
            
            query = """
            query getTestsFromSets($jql: String!) {
              getTestSets(jql: $jql, limit: 100) {
                results {
                  jira(fields: ["key"])
                  tests(limit: 100) {
                    results {
                      jira(fields: ["key"])
                      steps {
                        action
                        result
                      }
                    }
                  }
                }
              }
            }
            """
            variables = {"jql": jql}
            
            logger.info(f"Calling Xray GraphQL API with JQL: {jql}")
            try:
                resp = requests.post(gql_url, headers=headers, json={"query": query, "variables": variables}, timeout=30)
                
                if resp.status_code == 200:
                    result = resp.json()
                    if 'errors' in result:
                        logger.error(f"GraphQL errors: {result['errors']}")
                        return jsonify({"success": False, "error": "Xray GraphQL Error"}), 500
                        
                    data = result.get('data', {})
                    test_sets = data.get('getTestSets', {}).get('results', [])
                    
                    for ts in test_sets:
                        ts_key = ts.get('jira', {}).get('key')
                        test_results = ts.get('tests', {}).get('results', [])
                        for r in test_results:
                            t_key = r.get('jira', {}).get('key')
                            if not t_key: continue
                            
                            all_test_keys.add(t_key)
                            if t_key not in test_to_set_map:
                                test_to_set_map[t_key] = ts_key
                            
                            # Process steps
                            steps = r.get('steps', [])
                            if steps:
                                actions = []
                                results = []
                                for idx, s in enumerate(steps):
                                    actions.append(f"{idx+1}. {s.get('action', '')}")
                                    results.append(f"{idx+1}. {s.get('result', '')}")
                                
                                test_steps_map[t_key] = {
                                    'steps_text': "\n".join(actions),
                                    'expected_text': "\n".join(results)
                                }
                else:
                    logger.error(f"Xray GraphQL error (Status {resp.status_code}): {resp.text}")
                    return jsonify({"success": False, "error": f"Xray API Error ({resp.status_code})"}), 500
            except Exception as e:
                logger.error(f"Exception during Xray GraphQL call: {str(e)}")
                return jsonify({"success": False, "error": str(e)}), 500

        if not all_test_keys:
            logger.info("No test keys identified for the selected Test Sets.")
            return jsonify({"success": True, "tests": [], "count": 0})
            
        logger.info(f"Identifying details for {len(all_test_keys)} unique test keys from Jira")
            
        # Fetch details for all unique test keys from Jira
        jira_config = get_user_jira_config()
        if not jira_config:
            return jsonify({"success": False, "error": "Jira configuration required"}), 400
            
        base_url = jira_config['url'].rstrip('/')
        auth = (jira_config['email'], jira_config['api_token'])
        jira_headers = {"Accept": "application/json", "Content-Type": "application/json"}
        
        # Singular 'issue' is the correct endpoint for Jira Cloud bulkfetch
        bulk_url = f"{base_url}/rest/api/3/issue/bulkfetch"
        bulk_payload = {
            "issueIdsOrKeys": list(all_test_keys),
            "fields": ["summary", "description", "status", "issuelinks"]
        }
        
        logger.info(f"Calling Jira Bulk Fetch: {bulk_url}")
        bulk_resp = requests.post(bulk_url, json=bulk_payload, headers=jira_headers, auth=auth, timeout=30)
        
        if bulk_resp.status_code != 200:
            logger.error(f"Jira bulk fetch failed (Status {bulk_resp.status_code}): {bulk_resp.text}")
            return jsonify({"success": False, "error": f"Jira bulk fetch failed ({bulk_resp.status_code})"}), 500

        tests_details = []
        issues = bulk_resp.json().get('issues', [])
        logger.info(f"Successfully fetched {len(issues)} issues from Jira")

        for issue in issues:
            fields = issue.get('fields', {})
            key = issue['key']
            
            # Extract description and convert from ADF if necessary
            desc_val = fields.get('description')
            text_desc = adf_to_text(desc_val) if desc_val else ""
            
            # Get steps and expected result from our GraphQL map
            steps_info = test_steps_map.get(key, {})
            steps_text = steps_info.get('steps_text', "")
            expected_text = steps_info.get('expected_text', "See Jira for details")

            # Extract User Story (Requirement) from issuelinks
            user_story_key = None
            for link in fields.get('issuelinks', []):
                # Look for inwardIssue or outwardIssue that looks like a story/req
                linked_issue = link.get('inwardIssue') or link.get('outwardIssue')
                if linked_issue:
                    link_name = link.get('type', {}).get('name', '').lower()
                    # Common names: "Tests", "Tested by", "Relates", etc.
                    # We often want the one that is NOT a test set
                    if 'test' in link_name or 'cover' in link_name:
                        user_story_key = linked_issue.get('key')
                        break

            tests_details.append({
                'key': key,
                'summary': fields.get('summary'),
                'description': text_desc,
                'steps': steps_text,
                'expected': expected_text,
                'test_set': test_to_set_map.get(key),
                'user_story': user_story_key
            })
        
        return jsonify({
            "success": True, 
            "tests": tests_details, 
            "count": len(tests_details),
            "jira_url": base_url
        })
        
    except Exception as e:
        logger.error(f"Error fetching tests for sets: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/xray/import-tests', methods=['POST'])
def import_tests_to_xray():
    """Import test cases to Xray Cloud using bulk API"""
    try:
        data = request.json
        logger.info(f"=== XRAY IMPORT REQUEST ===")
        
        table_data = data.get('table_data', [])
        project_key = data.get('project_key')
        issue_key = data.get('issue_key')  # Extract issue key for linking
        test_set_key = data.get('test_set_key') # Extract test set key
        create_test_set_from_story = data.get('create_test_set_from_story', False)
        test_set_name = data.get('selected_issue_name') # Extract test set key
        
        logger.info(f"Table data rows: {len(table_data)}")
        logger.info(f"Project Key: {project_key}")
        logger.info(f"Test Set Name: {test_set_name}")

        jira_config = get_user_jira_config()
        if not jira_config:
            return jsonify({"error": "Jira configuration required"}), 400

        # Handle Auto-creation of Test Set
        if create_test_set_from_story and issue_key:
            # Prefer the name sent from the frontend (which is already parsed)
            story_summary = test_set_name
                
            if story_summary:
                logger.info(f"Creating Test Set with summary: {story_summary}")
                new_ts_key = create_jira_issue(project_key, story_summary, "Test Set", jira_config)
                if new_ts_key:
                    test_set_key = new_ts_key
                    logger.info(f"Successfully created new Test Set: {test_set_key}")
                else:
                    logger.warning("Failed to create Test Set (Check permissions). Proceeding with import...")
            else:
                logger.warning(f"Could not determine summary for {issue_key}. Proceeding with import...")

        if issue_key:
            logger.info(f"Linking to Issue: {issue_key}")
        if test_set_key:
            logger.info(f"Adding to Test Set: {test_set_key}")
        
        if not table_data or len(table_data) < 2:
            return jsonify({"error": "No valid table data provided"}), 400
        
        if not project_key:
            return jsonify({"error": "Project key is required"}), 400
        
        test_cases = parse_test_cases_table(table_data)
        logger.info(f"Parsed {len(test_cases)} test cases")
        
        if not test_cases:
            return jsonify({"error": "No valid test cases found in table data"}), 400
        
        success, result, error_msg = import_tests_to_xray_bulk(test_cases, project_key, issue_key=issue_key, test_set_key=test_set_key, poll_job=True)
        
        if success:
            # Handle linking if issue_key is provided
            if issue_key:
                created_keys = []
                # Check for sync result
                if isinstance(result, dict) and 'details' in result and isinstance(result['details'], list):
                     created_keys = [i.get('key') for i in result['details'] if i.get('key')]
                # Check for async job result
                elif isinstance(result, dict) and 'result' in result and isinstance(result['result'], dict) and 'issues' in result['result']:
                     created_keys = [i.get('key') for i in result['result']['issues'] if i.get('key')]
                # Check if result IS the job data (poll_import_job_status returns (True, job_data, None))
                elif isinstance(result, dict) and 'status' in result and 'result' in result:
                     if isinstance(result['result'], dict) and 'issues' in result['result']:
                          created_keys = [i.get('key') for i in result['result']['issues'] if i.get('key')]

                if created_keys:
                    logger.info(f"Linking {len(created_keys)} tests to {issue_key}")
                    link_success, link_msg = link_jira_issues(created_keys, issue_key)
                    
                    if link_success:
                        # Ensure result structure allows adding info
                        if 'result' not in result or not isinstance(result.get('result'), dict):
                             # If result is just a wrapper, we might need to adjust.
                             # But usually 'result' is the job data.
                             pass
                        
                        # Add linking info to response
                        if isinstance(result, dict):
                            if 'result' in result and isinstance(result['result'], dict):
                                result['result']['linked_count'] = len(created_keys)
                                result['result']['linked_issue'] = issue_key
                            else:
                                result['linked_count'] = len(created_keys)
                                result['linked_issue'] = issue_key
                    else:
                        if isinstance(result, dict):
                            if 'result' in result and isinstance(result['result'], dict):
                                result['result']['linking_error'] = link_msg
                            else:
                                result['linking_error'] = link_msg

                # Handle Test Set Association
                if test_set_key and created_keys:
                    logger.info(f"Adding {len(created_keys)} tests to Test Set: {test_set_key}")
                    ts_success, ts_msg = add_tests_to_xray_test_set(created_keys, test_set_key)
                    
                    if ts_success:
                        if isinstance(result, dict):
                            if 'result' in result and isinstance(result['result'], dict):
                                result['result']['test_set_added'] = True
                                result['result']['test_set_key'] = test_set_key
                            else:
                                result['test_set_added'] = True
                                result['test_set_key'] = test_set_key
                    else:
                        if isinstance(result, dict):
                            if 'result' in result and isinstance(result['result'], dict):
                                result['result']['test_set_error'] = ts_msg
                            else:
                                result['test_set_error'] = ts_msg

                else:
                    logger.warning("No created test keys found to link")

            return jsonify({
                "success": True,
                "message": f"Successfully submitted {len(test_cases)} test cases to Xray",
                "result": result,
                "note": "Check Xray dashboard for import status"
            })
        else:
            return jsonify({
                "success": False,
                "error": error_msg
            }), 400
        
    except Exception as e:
        logger.error(f"Error in import_tests_to_xray: {str(e)}")
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

# MCB Server Endpoints
@app.route('/mcb/test-connection', methods=['GET'])
def test_mcb_connection():
    try:
        if MOCK_MCB:
            return jsonify({
                "success": True,
                "message": "MCB mock mode enabled",
                "version": "mock-local"
            })
        test_url = f"{MCB_SERVER_URL}/health"
        try:
            response = requests.get(test_url, timeout=10)
            if response.status_code == 200:
                return jsonify({
                    "success": True,
                    "message": "MCB Server connection successful",
                    "version": response.json().get('version', 'unknown')
                })
            return jsonify({
                "success": False,
                "error": f"MCB Server returned status {response.status_code}"
            }), 500
        except requests.exceptions.RequestException as e:
            return jsonify({
                "success": False,
                "error": f"MCB Server not reachable: {str(e)}"
            }), 500
    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"MCB connection test failed: {str(e)}"
        }), 500

@app.route('/mcb/execute-workflow', methods=['POST'])
def execute_mcb_workflow():
    try:
        if not session.get('authenticated'):
            return jsonify({"success": False, "error": "Authentication required"}), 401
        data = request.json or {}
        description = (data.get('description') or '').strip()
        context = data.get('context') or {}
        if not description:
            return jsonify({"success": False, "error": "description is required"}), 400
        if MOCK_MCB:
            try:
                code_prompt = (
                    "You are a test automation engineer. Using the following natural language test workflow, "
                    "generate a complete Playwright test written in TypeScript using the @playwright/test runner. "
                    "Use a single describe block and a single test that implements the workflow. Respond with code only.\n\n"
                    f"Workflow description:\n{description}"
                )
                code = call_mistral_model(code_prompt)
            except Exception as e:
                logger.error(f"Error generating Playwright code in mock MCB: {str(e)}")
                return jsonify({"success": False, "error": "Failed to generate Playwright test"}), 500
            filename = f"nl_workflow_{uuid.uuid4().hex[:8]}.spec.ts"
            output_path = os.path.join("/tmp", filename)
            try:
                os.makedirs("/tmp", exist_ok=True)
            except Exception:
                pass
            try:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(code)
            except Exception as e:
                logger.error(f"Failed to write workflow spec file: {str(e)}")
                return jsonify({"success": False, "error": "Failed to write spec file"}), 500
            run_id = start_playwright_run(output_path)
            workflow = {
                "mode": "mock",
                "description": description,
                "context": context,
                "execution": {
                    "tool": "playwright",
                    "status": "running",
                    "run_id": run_id,
                    "spec_file": filename
                }
            }
            return jsonify({"success": True, "workflow": workflow})
        payload = {
            "description": description,
            "context": context,
            "user": session.get('username')
        }
        url = f"{MCB_SERVER_URL.rstrip('/')}/workflow/execute"
        try:
            response = requests.post(url, json=payload, timeout=60)
        except requests.exceptions.RequestException as e:
            logger.error(f"MCB request failed: {str(e)}")
            return jsonify({"success": False, "error": f"MCB request failed: {str(e)}"}), 502
        try:
            body = response.json()
        except ValueError:
            body = {"raw": response.text}
        if response.status_code >= 400:
            return jsonify({
                "success": False,
                "status": response.status_code,
                "error": body
            }), 502
        return jsonify({"success": True, "workflow": body})
    except Exception as e:
        logger.error(f"Error in execute_mcb_workflow: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

@app.route('/mcb/generate-playwright-test', methods=['POST'])
def generate_playwright_test():
    try:
        if not session.get('authenticated'):
            return jsonify({"success": False, "error": "Authentication required"}), 401
        data = request.json or {}
        description = (data.get('description') or '').strip()
        context = data.get('context') or {}
        if not description:
            return jsonify({"success": False, "error": "description is required"}), 400
        if MOCK_MCB:
            code_prompt = (
                "You are a test automation engineer. Using the following natural language test request, "
                "generate a complete Playwright test written in TypeScript using the "
                "@playwright/test runner. Use a single describe block and a single test. "
                "Respond with code only.\n\n"
                f"Request:\n{description}"
            )
            code = call_mistral_model(code_prompt)
            filename = f"playwright_test_{uuid.uuid4().hex[:8]}.spec.ts"
            output_path = os.path.join("/tmp", filename)
            try:
                os.makedirs("/tmp", exist_ok=True)
            except Exception:
                pass
            try:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(code)
            except Exception as e:
                logger.error(f"Failed to write Playwright spec file: {str(e)}")
                return jsonify({"success": False, "error": "Failed to write spec file"}), 500
            return jsonify({
                "success": True,
                "code": code,
                "download": f"/download/{filename}"
            })
        payload = {
            "description": description,
            "context": context,
            "user": session.get('username')
        }
        url = f"{MCB_SERVER_URL.rstrip('/')}/playwright/generate-test"
        try:
            response = requests.post(url, json=payload, timeout=60)
        except requests.exceptions.RequestException as e:
            logger.error(f"MCB Playwright generation failed: {str(e)}")
            return jsonify({"success": False, "error": f"MCB Playwright request failed: {str(e)}"}), 502
        try:
            body = response.json()
        except ValueError:
            body = {"raw": response.text}
        if response.status_code >= 400:
            return jsonify({
                "success": False,
                "status": response.status_code,
                "error": body
            }), 502
        return jsonify(body)
    except Exception as e:
        logger.error(f"Error in generate_playwright_test: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

@app.route('/mcb/run-playwright-test', methods=['POST'])
def run_playwright_test():
    try:
        if not session.get('authenticated'):
            return jsonify({"success": False, "error": "Authentication required"}), 401
        data = request.json or {}
        filename = (data.get('filename') or '').strip()
        if not filename:
            return jsonify({"success": False, "error": "filename is required"}), 400
        spec_path = os.path.join("/tmp", filename)
        if not os.path.exists(spec_path):
            return jsonify({"success": False, "error": "spec file not found"}), 400
        run_id = start_playwright_run(spec_path)
        return jsonify({"success": True, "run_id": run_id})
    except Exception as e:
        logger.error(f"Error in run_playwright_test: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

@app.route('/mcb/playwright-run-status', methods=['GET'])
def playwright_run_status():
    try:
        run_id = (request.args.get('run_id') or '').strip()
        if not run_id:
            return jsonify({"success": False, "error": "run_id is required"}), 400
        run = PLAYWRIGHT_RUNS.get(run_id)
        if not run:
            return jsonify({"success": False, "error": "run not found"}), 404
        return jsonify({
            "success": True,
            "status": run.get("status"),
            "logs": run.get("logs", []),
            "returncode": run.get("returncode")
        })
    except Exception as e:
        logger.error(f"Error in playwright_run_status: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "error": f"Internal server error: {str(e)}"}), 500

# AI Generation Endpoints
@app.route('/generate-test-plan', methods=['POST'])
def generate_test_plan():
    try:
        data = request.json
        job_scope = data.get('job_scope')
        template_path = os.path.join(os.path.dirname(__file__), "templates", "test_plan_template.docx")
        # output_path = os.path.join("/tmp", "Updated_Test_Plan.docx")
        output_path = os.path.join("templates", "Updated_Test_Plan.docx")
        if not job_scope:
            return jsonify({"error": "job_scope is required"}), 400
        
        if not os.path.exists(template_path):
            return jsonify({"error": f"Template file not found at {template_path}"}), 400
            
        sections = extract_sections_from_template(template_path)
        filled_sections = fill_sections_with_scope(sections, job_scope)
        result_path = update_template_with_filled_sections(template_path, filled_sections, output_path)
        
        logger.info(f"Test plan generated successfully at: {result_path}")
        return jsonify({"message": "Test plan generated", "output_path": result_path})
        
    except Exception as e:
        logger.error(f"Error generating test plan: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/generate-test-cases', methods=['POST'])
def generate_test_cases():
    try:
        data = request.json
        user_story = data.get('user_story')

        if not user_story:
            return jsonify({"error": "user_story is required"}), 400

        prompt = f"{user_story}, from the given user story create a table of test cases (Test Case ID, Description (starts with validate that....), Steps, Expected Behavior) with corner cases, give me the table only"
        result = call_mistral_model(prompt)
        
        logger.info("Test cases generated successfully")
        return jsonify({"message": "Test Cases generated", "output": result})
        
    except Exception as e:
        logger.error(f"Error generating test cases: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/review-test-cases', methods=['POST'])
def review_test_cases():
    data = request.json
    test_cases = data.get('test_cases')
    print(test_cases)

    if not test_cases:
        return jsonify({"error": "test_cases is required"}), 400

    prompt = f"""{test_cases}, from the given test cases, review each test case based on the following analysis:
    1. Test Case ID: Each test case should have a unique identifier.
    2. Test Case name is self explanatory:	The name of the test case clearly indicates its purpose and easy to understand what  module is the test subject.
    3. Test Objective:	Clearly states what the test is intended to verify.
    4. Functionality coverage:	All relevant functionalities are covered.
    5. Test Data and Pre-Requisites Defined:	Test Data and Pre-Requisites for execution are clearly specified.
    6. Correct Sequence of Steps:	Steps/Actions should state very clearly the sequence of actions to be carried out on the system by the user and easy to follow 
    7. Expected Results Defined: 
        - Expected results are simple and clearly stated
        - Expected Results should clearly state how the system should respond to the user actions given in each step/action.    
        - Ensure that too many things are not included to be verified under one expected output. 
        - Ensure that separate cases are written for multiple verifications of the application’s behavior.
        - Vague statements like “Appropriate message/value/screen” etc, should not be part of expected result. Every detail should be clearly spelt out."
    8. Priority Level:	Indicates the importance/urgency of the test.
    9. Testability:	The test case is practical and feasible to execute.(explicit and implicit)
    10. Accurate:	The test case provides correct and valid verification steps and It is obvious what the test is trying to test.
    11. Economical:	Test case avoids unnecessary steps and uses minimal resources.
    12. Redundancy:	No duplicate test cases or steps are present.
    13. Replication:	Test case can be re-executed with consistent results.
    14. Grammatical & Spelling: 	All statements are free from grammatical and spelling mistakes.
    15. Simple Language:	Test case uses clear and simple langauge.
    16. Reference File Attached:	Reference files are attached and accurate.
    17. Correct Directory Placement:	Test case is placed in the correct directory.
    18. Configuration Information: 	"Specifies the configuration information such as :
                                    *Environment Details, 
                                    *Test Data, 
                                    *Test pre-requisite, 
                                    *Security Access( if exist)"
    19. Requirement to Test Condition Mapping: All requirements are mapped to test conditions, both explicit and implicit, been converted into Test conditions.
    20. Boundary/Special/Invalid: 	"These values are properly identified and tested, Boundary values, Special values and Invalid values been identified and included in the Test cases."
    21. Negative Scenarios: 	Negative test conditions are included.
    22. Reviewer Name/Date:	For QA reviewer to confirm test case validity.
    ,
    Create a table with the analysis factor name, Complies (Yes/No), and comments (if Complies is No, with a reference to test cases (e.g., test case id)), then Givve me the table only.
    """
    result = call_mistral_model(prompt)
    print(result)
    print({"message": "Test Cases Reviewed", "output": result})
    return jsonify({"message": "Test Cases Reviewed", "output": result})

@app.route('/generate-bug-report', methods=['POST'])
def generate_bug_report():
    try:
        data = request.json
        bugs = data.get('bugs')

        if not bugs:
            return jsonify({"error": "bugs is required"}), 400

        prompt = f"{bugs}, from the given bugs, create a bug report as a table with all attributes (Bug ID,Description,Steps to Reproduce,Expected Behavior,Actual Behavior,Priority,Severity,Status), give me the table only"
        result = call_mistral_model(prompt)
        
        logger.info("Bug report generated successfully")
        return jsonify({"message": "Bug report generated", "output": result})
        
    except Exception as e:
        logger.error(f"Error generating bug report: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/generate-data', methods=['POST'])
def generate_data():
    try:
        data = request.json
        input_text = data.get('input')

        if not input_text:
            return jsonify({"error": "input is required"}), 400

        prompt = f"""You are a dummy data generator.  
I will provide you with an input description.  
You must always produce only a table as output, with exactly the column names and format I specify in the description.  

Rules:  
- Output must be a table only (no explanations, no extra text).  
- Column headers must exactly match what I specify.  
- Each row must contain one complete and realistic data record.  
- Number of rows = as specified in my input.  
- Data must look realistic (not just random gibberish).  
- Ensure uniqueness where it makes sense (e.g., IDs, emails).  

Example inputs:  
1. "10 questions about employees, their salaries and departments → table with 2 columns (question #, question)"  
2. "emails and passwords → table with 2 columns (email, password), 15 rows"  
3. "product codes and prices → table with 2 columns (product_code, price), 20 rows"  

Now wait for my input and generate the table accordingly. 

input description: {input_text}"""
        
        result = call_mistral_model(prompt)
        logger.info("Dummy data generated successfully")
        return jsonify({"message": "Dummy data generated", "output": result})
        
    except Exception as e:
        logger.error(f"Error generating data: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/analyse-req', methods=['POST'])
def analyse_requirements():
    data = request.json
    user_story = data.get('user_story')
    print(user_story)

    if not user_story:
        return jsonify({"error": "user_story is required"}), 400

    prompt = f"""{user_story}
            From the given user stories/requirements, 
            Create a table of comprehensive analysis for the given requirements/user stories, including only (#, Comments/Questions/Gaps (e.g. comprehensive questions if there are any ambiguous), Answers/Assumptions), 
            Give me the created table only"""
    result = call_mistral_model(prompt)
    print(result)
    print({"message": "Requirement Analysis generated", "output": result})
    return jsonify({"message": "Requirement Analysis generated", "output": result})

@app.route('/generate-hls', methods=['POST'])
def generate_hls():
    data = request.json
    user_story = data.get('user_story')
    print(user_story)

    if not user_story:
        return jsonify({"error": "user_story is required"}), 400

    prompt = f"""{user_story}
            From the given user stories/requirements, 
            Create a table of comprehensive high-level scenarios that analyse the given requirements/user stories (User Story ID, User Story Name, Scenario ID, Module, Scenario Priority, Scenario Description), 
            Give me the created table only"""
    result = call_mistral_model(prompt)
    print(result)
    print({"message": "High Level Scenarios generated", "output": result})
    return jsonify({"message": "High Level Scenarios generated", "output": result})

@app.route('/generate-data-quality-tests', methods=['POST'])
def generate_data_quality_tests():
    try:
        data = request.json
        input_schema = data.get('schema')

        if not input_schema:
            return jsonify({"error": "schema is required"}), 400

        prompt = f"""You are a mySQL data quality and testing expert.
    Given a table schema, table name, and list of columns with their data types, generate a comprehensive SQL validation script that covers both schema tests and data quality tests, returning a unified result set with the following columns:
    Check_Type | Dimension | Rule | Expectation | Failed_Records

    Check_Type (e.g., Schema, Data,...)
    Dimension must be in (e.g., Completeness, Integrity, Uniqueness, Validity, Consistency,...)
    Rule (e.g. Invalid mental_health_rating, Attendance<50 but Exam>95,...)
    Expectation (e.g. Diet quality must be Poor, Average, or Good, If sleep_hours<3, mental_health_rating should not exceed 9,...)
    Failed_Records (e.g. number of failed records)

    Your task:

    Schema Tests:

    - Row count test → Verify expected record count (use a placeholder number like 2000).

    - Column count test → Verify number of columns matches the provided schema.

    - Column names and types test → Verify each column matches its expected data type.

    Data Quality Tests:

    - Uniqueness: Check if key columns (like IDs) are unique.

    - Completeness: Check for nulls in critical fields (IDs, required attributes, metrics).

    - Validity: Based on data type and semantic meaning, create range and categorical validation rules:

    - For numeric columns: specify realistic ranges (e.g., 0–100 for percentages or scores).

    - For categorical columns: specify allowed values (e.g., Gender: Male/Female/Other).

    - Consistency: Create logical dependency rules (e.g., if attendance < 50, exam_score < 95).

    - Use UNION ALL between each test block.

    - Include descriptive comments (e.g., -- =======================).

    - Use the given schema and table name consistently.

    Input Example:
    
    {input_schema}


    Output Format Should be the script only, Return a single SQL script that includes:

    - Clear comment sections (-- SCHEMA TESTS, -- DATA QUALITY TESTS)

    - A combined query using UNION ALL

    - Proper CASE logic for validation rules

    - Friendly readable expectations (e.g., 'Attendance must be between 0 and 100')

    The final output should look similar in structure to the following:
    -- =======================
    -- SCHEMA TESTS
    -- =======================
    SELECT ...
    UNION ALL
    ...
    -- =======================
    -- DATA QUALITY TESTS
    -- =======================
    SELECT ...
    """

        result = call_mistral_model(prompt)
        logger.info("Data quality tests generated successfully")
        return jsonify({"message": "Data Quality Tests generated", "output": result})
        
    except Exception as e:
        logger.error(f"Error generating data quality tests: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/chat-about-testcase', methods=['POST'])
def chat_about_testcase():
    """Chat with AI about a specific test case"""
    try:
        data = request.json
        test_case = data.get('test_case')
        user_story = data.get('user_story')
        message = data.get('message')
        history = data.get('history', [])

        if not message:
            return jsonify({"error": "Message is required"}), 400

        # Build context
        context = f"""
You are a QA Assistant helping a user with a specific test case.
        
User Story:
{user_story or 'Not provided'}

Test Case Details:
ID: {test_case.get('id', 'N/A')}
Description: {test_case.get('description', 'N/A')}
Steps: {test_case.get('steps', 'N/A')}
Expected Result: {test_case.get('expected_result', 'N/A')}

Chat History:
"""
        for msg in history:
            role = "User" if msg.get('sender') == 'user' else "AI"
            context += f"{role}: {msg.get('text')}\n"

        prompt = f"""{context}

User: {message}

Answer as the AI QA Assistant. Be helpful, concise, and focused on testing best practices.
"""
        
        reply = call_mistral_model(prompt)
        return jsonify({"reply": reply})

    except Exception as e:
        logger.error(f"Error in chat_about_testcase: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Authentication Configuration
def require_auth(f):
    """Decorator to require authentication"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('authenticated'):
            return jsonify({"success": False, "error": "Authentication required"}), 401
        return f(*args, **kwargs)
    return decorated_function

# Authentication Routes
@app.route("/login", methods=["GET"])
def login_page():
    """Show login page if not authenticated"""
    if session.get('authenticated'):
        return redirect('/')
    return render_template("login.html")

@app.route("/login", methods=["POST"])
def login():
    """Handle login request - authenticate with Jira"""
    try:
        data = request.json
        jira_url = data.get('jira_url', '').strip().rstrip('/')
        jira_email = data.get('jira_email', '').strip()
        jira_api_token = data.get('jira_api_token', '').strip()
        
        if not jira_url or not jira_email or not jira_api_token:
            return jsonify({"success": False, "error": "Please provide Jira URL, email, and API token"}), 400
        
        # Test Jira connection
        jira_config = {
            'url': jira_url,
            'email': jira_email,
            'api_token': jira_api_token
        }
        
        success, message, user_data = test_jira_connection(jira_config)
        
        if success:
            # Store Jira config in session
            session['authenticated'] = True
            session['jira_config'] = jira_config
            session['jira_user'] = user_data.get('displayName', jira_email) if user_data else jira_email
            session['username'] = jira_email
            
            logger.info(f"User {jira_email} logged in successfully with Jira")
            return jsonify({
                "success": True, 
                "message": "Login successful",
                "user": user_data.get('displayName', jira_email) if user_data else jira_email
            })
        else:
            logger.warning(f"Failed Jira login attempt for email: {jira_email}")
            return jsonify({"success": False, "error": message or "Invalid Jira credentials"}), 401
    except Exception as e:
        logger.error(f"Login error: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "error": f"Login failed: {str(e)}"}), 500

@app.route("/logout", methods=["POST"])
def logout():
    """Handle logout request"""
    session.clear()
    logger.info("User logged out")
    return jsonify({"success": True, "message": "Logged out successfully"})

@app.route("/check-auth", methods=["GET"])
def check_auth():
    """Check if user is authenticated"""
    return jsonify({
        "authenticated": session.get('authenticated', False),
        "username": session.get('username')
    })

# Serve home page (protected)
@app.route("/", methods=["GET"])
def home():
    """Home page - requires authentication"""
    if not session.get('authenticated'):
        return redirect('/login')
    return render_template("index.html")

@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"500 Error: {str(error)}")
    return jsonify({"error": "Internal server error"}), 500

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({"error": "Endpoint not found"}), 404

if __name__ == "__main__":
    templates_dir = os.path.join(os.path.dirname(__file__), "templates")
    if not os.path.exists(templates_dir):
        os.makedirs(templates_dir)
        logger.info(f"Created templates directory: {templates_dir}")
    
    template_path = os.path.join(templates_dir, "test_plan_template.docx")
    if not os.path.exists(template_path):
        logger.warning(f"Template file not found: {template_path}")
        logger.warning("Please make sure the test_plan_template.docx file exists in the templates directory")
    
    app.run(debug=True, port=8084)
